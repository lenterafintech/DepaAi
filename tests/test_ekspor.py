"""Uji ekspor laporan ke berbagai format berkas dan sintaks yang dapat dijalankan ulang."""

from __future__ import annotations

import io
import json
import zipfile
from pathlib import Path

import pandas as pd
import pytest

from nalardata import ekspor, sintaks
from nalardata import narrative as nr

ROOT = Path(__file__).resolve().parents[1]
NUMERIK = [
    "usia",
    "lama_usaha_tahun",
    "pendapatan_bulanan",
    "saldo_tabungan",
    "rasio_utang_pendapatan",
    "skor_kredit",
    "jumlah_keterlambatan",
]


@pytest.fixture(scope="module")
def laporan() -> nr.Laporan:
    df = pd.read_csv(ROOT / "data" / "contoh_data_nasabah.csv")
    konfig = nr.Konfigurasi(
        variabel=NUMERIK,
        nama_data="contoh_data_nasabah.csv",
        target_numerik="skor_kredit",
        prediktor=["rasio_utang_pendapatan", "jumlah_keterlambatan", "pendapatan_bulanan"],
        target_biner="gagal_bayar",
        prediktor_biner=["rasio_utang_pendapatan", "skor_kredit"],
        kelompok="segmen_usaha",
        gugus_x=["pendapatan_bulanan", "saldo_tabungan"],
        gugus_y=["skor_kredit", "jumlah_keterlambatan"],
    )
    _, lap = nr.analisis_dan_laporan(df, konfig)
    return lap


def test_konfigurasi_ikut_tersimpan(laporan):
    assert laporan.konfig is not None
    assert laporan.konfig.target_numerik == "skor_kredit"


@pytest.mark.parametrize("kode", sorted(ekspor.FORMAT))
@pytest.mark.parametrize("lengkap", [False, True])
def test_setiap_format_menghasilkan_berkas(laporan, kode, lengkap):
    isi = ekspor.bangun(laporan, kode, "akademik", lengkap)
    assert isinstance(isi, bytes)
    assert len(isi) > 500


def test_format_ditolak_bila_tak_dikenal(laporan):
    with pytest.raises(ValueError, match="tidak dikenal"):
        ekspor.bangun(laporan, "xls", "akademik")
    with pytest.raises(ValueError, match="Pembaca"):
        ekspor.bangun(laporan, "html", "manajer")


def test_docx_terbaca_kembali(laporan):
    from docx import Document

    dok = Document(io.BytesIO(ekspor.ke_docx(laporan, "akademik", lengkap=True)))
    teks = "\n".join(p.text for p in dok.paragraphs)
    assert "Laporan Lengkap Analisis Multivariat" in teks
    assert laporan.headline in teks
    assert dok.tables, "laporan lengkap harus memuat tabel"


def test_pptx_terbaca_kembali(laporan):
    from pptx import Presentation

    presentasi = Presentation(io.BytesIO(ekspor.ke_pptx(laporan, "eksekutif")))
    assert len(presentasi.slides) >= 3


def test_pdf_terbaca_kembali(laporan):
    from pypdf import PdfReader

    baca = PdfReader(io.BytesIO(ekspor.ke_pdf(laporan, "akademik", lengkap=True)))
    assert len(baca.pages) >= 2
    assert "Laporan Lengkap" in baca.pages[0].extract_text()


def test_xlsx_berisi_lembar_per_bagian(laporan):
    lembar = pd.read_excel(
        io.BytesIO(ekspor.ke_xlsx(laporan, "akademik", lengkap=True)), sheet_name=None
    )
    assert "Keterangan" in lembar
    # Tiap tabel hasil menjadi lembar tersendiri, sehingga angkanya siap diolah ulang.
    assert len(lembar) > 3
    # Nama lembar Excel dibatasi 31 karakter dan harus unik.
    assert all(len(n) <= 31 for n in lembar)
    assert len(set(n.lower() for n in lembar)) == len(lembar)


def test_json_terstruktur(laporan):
    isi = json.loads(ekspor.ke_json(laporan, "profesional", lengkap=True))
    assert isi["judul"].endswith(nr.AUDIENCE_LABELS["profesional"])
    assert isi["meta"]
    jenis = {b["jenis"] for b in isi["blok"]}
    assert {"judul", "meta", "subjudul", "tabel"} <= jenis
    # Tabel ikut terbawa sebagai baris terstruktur, bukan sekadar teks.
    tabel = [b for b in isi["blok"] if b["tabel"]]
    assert tabel and isinstance(tabel[0]["tabel"], list)


def test_html_ringkas_hanya_satu_register(laporan):
    satu = ekspor.ke_html(laporan, "eksekutif").decode("utf-8")
    semua = ekspor.ke_html(laporan, "eksekutif", lengkap=True).decode("utf-8")
    assert len(semua) > len(satu)
    assert "<html" in satu.lower() or "<!doctype" in satu.lower()


def test_zip_memuat_seluruh_format_dan_sintaks(laporan):
    arsip = zipfile.ZipFile(io.BytesIO(ekspor.ke_zip(laporan, "akademik", lengkap=True)))
    nama = arsip.namelist()
    for akhiran in (".html", ".md", ".docx", ".xlsx", ".pdf", ".json"):
        assert any(n.endswith(akhiran) for n in nama), akhiran
    assert "sintaks/analisis.py" in nama
    assert "sintaks/analisis.R" in nama
    assert any(n.startswith("tabel/") and n.endswith(".csv") for n in nama)


def test_nama_berkas_bersih(laporan):
    assert ekspor.nama_berkas(laporan, "docx", "akademik").endswith(
        "_ringkasan_akademik.docx"
    )
    assert ekspor.nama_berkas(laporan, "pdf", "akademik", True).endswith(
        "_laporan_lengkap_akademik.pdf"
    )
    # Sintaks tidak bergantung pada register pembaca.
    assert ekspor.nama_berkas(laporan, "py", "eksekutif") == ekspor.nama_berkas(
        laporan, "py", "akademik", True
    )


def test_nama_lembar_unik_dan_pendek():
    dipakai: set[str] = set()
    a = ekspor._nama_lembar("Tabel yang judulnya sangat panjang sekali", dipakai)
    b = ekspor._nama_lembar("Tabel yang judulnya sangat panjang sekali", dipakai)
    assert len(a) <= 31 and len(b) <= 31
    assert a != b
    assert ekspor._nama_lembar("a[b]c:d*e?f/g", dipakai) == "abcdefg"


def test_laporan_lengkap_tetap_satu_register(laporan):
    """Laporan lengkap adalah dokumen yang lebih dalam, bukan tiga ringkasan digabung."""
    lengkap = ekspor.susun_blok(laporan, "eksekutif", lengkap=True)
    teks_gabung = " ".join(b.teks for b in lengkap)
    # Label register lain tidak boleh muncul sebagai tajuk bagian.
    for lain in ("akademik", "profesional"):
        assert nr.AUDIENCE_LABELS[lain] not in teks_gabung
    assert nr.AUDIENCE_LABELS["eksekutif"] in teks_gabung

    # Uraian tiap temuan memakai register yang diminta, bukan register lain.
    uraian = {b.teks for b in lengkap if b.jenis == "paragraf"}
    assert laporan.temuan[0].eksekutif in uraian
    assert laporan.temuan[0].akademik not in uraian


def test_lengkap_lebih_dalam_daripada_ringkasan(laporan):
    for pembaca in nr.AUDIENCES:
        ringkas = ekspor.susun_blok(laporan, pembaca, lengkap=False)
        lengkap = ekspor.susun_blok(laporan, pembaca, lengkap=True)
        assert len(lengkap) > len(ringkas)
        # Tabel hasil dan rujukan hanya muncul pada laporan lengkap.
        assert sum(1 for b in lengkap if b.jenis == "tabel") > sum(
            1 for b in ringkas if b.jenis == "tabel"
        )


def test_susun_blok_menolak_pembaca_asing(laporan):
    with pytest.raises(ValueError, match="tidak dikenal"):
        ekspor.susun_blok(laporan, "manajer")


# --------------------------------------------------------------------------- #
# Sintaks yang dapat dijalankan ulang
# --------------------------------------------------------------------------- #


def test_sintaks_python_valid_secara_sintaksis(laporan):
    skrip = sintaks.sintaks_python(laporan.konfig)
    compile(skrip, "sintaks_analisis.py", "exec")
    assert "sm.OLS" in skrip
    assert "sm.Logit" in skrip
    assert "MANOVA" in skrip
    assert laporan.konfig.target_numerik in skrip


def test_sintaks_r_memuat_langkah_yang_dipilih(laporan):
    skrip = sintaks.sintaks_r(laporan.konfig)
    assert "lm(" in skrip
    assert "glm(" in skrip
    assert "manova(" in skrip
    assert laporan.konfig.kelompok in skrip


def test_sintaks_menyesuaikan_konfigurasi_minimal():
    konfig = nr.Konfigurasi(variabel=["a", "b"], nama_data="d.csv")
    py = sintaks.sintaks_python(konfig)
    compile(py, "d.py", "exec")
    assert "sm.OLS" not in py
    assert "MANOVA" not in py
    assert "lm(" not in sintaks.sintaks_r(konfig)


def test_sintaks_tanpa_konfigurasi_menjelaskan_sebabnya():
    teks = sintaks.bangkitkan(None, "py")
    assert "tidak dapat dibangkitkan" in teks
    with pytest.raises(ValueError):
        sintaks.bangkitkan(nr.Konfigurasi(variabel=["a"]), "sas")


def test_nama_variabel_dikutip_dengan_aman():
    konfig = nr.Konfigurasi(variabel=['aneh"nama'], nama_data='data"ku.csv')
    compile(sintaks.sintaks_python(konfig), "d.py", "exec")


# --------------------------------------------------------------------------- #
# Sintaks SPSS, AMOS, dan Mplus
# --------------------------------------------------------------------------- #


def test_sintaks_spss_memuat_perintah_yang_setara(laporan):
    skrip = sintaks.sintaks_spss(laporan.konfig)
    # Perintah pokok SPSS untuk analisis yang dipilih pada fixture.
    for perintah in ("GET DATA", "DESCRIPTIVES", "CORRELATIONS", "FACTOR", "REGRESSION"):
        assert perintah in skrip, perintah
    assert "LOGISTIC REGRESSION" in skrip  # target biner ada pada konfigurasi
    assert "GLM" in skrip and "DISCRIMINANT" in skrip  # kelompok ada
    assert laporan.konfig.target_numerik in skrip
    # Sintaks SPSS berakhir dengan RESTORE agar pengaturan sesi dikembalikan.
    assert skrip.rstrip().endswith("RESTORE.")


def test_sintaks_spss_menyesuaikan_konfigurasi_minimal():
    konfig = nr.Konfigurasi(variabel=["a", "b"], nama_data="d.csv")
    skrip = sintaks.sintaks_spss(konfig)
    assert "DESCRIPTIVES" in skrip
    assert "REGRESSION" not in skrip
    assert "GLM" not in skrip


def test_sintaks_amos_menyebut_jalur_dan_langkahnya(laporan):
    teks = sintaks.sintaks_amos(laporan.konfig)
    assert "VARIABEL TERAMATI" in teks
    assert "JALUR STRUKTURAL" in teks
    for prediktor in laporan.konfig.prediktor:
        assert f"{prediktor} -> {laporan.konfig.target_numerik}" in teks


def test_nama_mplus_unik_meski_awalannya_sama():
    """Pemotongan lugas akan menabrakkan nama; pemetaan harus tetap unik."""
    nama = ["pendapatan_bulanan", "pendapatan_tahunan", "pendapatan_harian", "usia"]
    peta = sintaks.nama_mplus(nama)
    assert len(set(peta.values())) == len(nama)
    assert all(len(pendek) <= 8 for pendek in peta.values())
    assert peta["usia"] == "usia"  # nama pendek tidak diubah


def test_sintaks_mplus_memakai_nama_yang_dipetakan(laporan):
    berkas = sintaks.sintaks_mplus(laporan.konfig)
    assert berkas.startswith("TITLE:")
    for bagian in ("DATA:", "VARIABLE:", "ANALYSIS:", "MODEL:", "OUTPUT:"):
        assert bagian in berkas, bagian
    assert "ESTIMATOR = MLR" in berkas
    # Seluruh nama pada berkas Mplus tidak boleh melewati 8 karakter.
    baris_model = [b for b in berkas.splitlines() if " ON " in b][0]
    for kata in baris_model.replace(";", "").split():
        if kata != "ON":
            assert len(kata) <= 8, kata


def test_bahasa_sintaks_tak_dikenal_ditolak(laporan):
    with pytest.raises(ValueError, match="tidak dikenal"):
        sintaks.bangkitkan(laporan.konfig, "sas")


@pytest.mark.parametrize("kode", ["spss", "amos", "mplus"])
def test_sintaks_baru_tersedia_sebagai_format_ekspor(laporan, kode):
    isi = ekspor.bangun(laporan, kode)
    assert len(isi) > 300
    # Sintaks tidak bergantung pada register pembaca.
    assert isi == ekspor.bangun(laporan, kode, "profesional", lengkap=True)
    assert ekspor.nama_berkas(laporan, kode).endswith(ekspor.FORMAT[kode].ekstensi)


def test_paket_zip_memuat_seluruh_sintaks_dan_petunjuknya(laporan):
    arsip = zipfile.ZipFile(io.BytesIO(ekspor.bangun(laporan, "zip", "akademik", True)))
    nama = arsip.namelist()
    for berkas in (
        "sintaks/analisis.py",
        "sintaks/analisis.R",
        "sintaks/analisis.sps",
        "sintaks/model_amos.txt",
        "sintaks/analisis.inp",
        "sintaks/BACA_DULU.txt",
    ):
        assert berkas in nama, berkas
    # Petunjuk harus menjelaskan bahwa data.csv tidak ikut dan cara mendapatkannya.
    petunjuk = arsip.read("sintaks/BACA_DULU.txt").decode("utf-8")
    assert "TIDAK disertakan" in petunjuk
    assert ".nalardata" in petunjuk
