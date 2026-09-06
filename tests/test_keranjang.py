"""Uji keranjang hasil dan ekspornya ke berkas."""

from __future__ import annotations

import io
import json
import zipfile

import pandas as pd
import pytest

from nalardata import ekspor
from nalardata import keranjang as kr


@pytest.fixture
def tabel() -> pd.DataFrame:
    return pd.DataFrame({"Variabel": ["x1", "x2"], "B": [0.2423, -0.0097]})


@pytest.fixture
def isi(tabel) -> kr.Keranjang:
    k = kr.Keranjang(judul="Laporan Uji", peneliti="Peneliti Satu")
    k.tambah_tabel("Regresi linear", "Koefisien", tabel, tanda_data="data#1")
    k.tambah_teks("Regresi linear", "Cara membaca", "Koefisien B menunjukkan perubahan.")
    k.tambah_tabel(
        "CFA", "Muatan", pd.DataFrame({"Butir": ["x1"], "Muatan": [0.81]}), tanda_data="data#1"
    )
    return k


# --------------------------------------------------------------------------- #
# Penyimpanan
# --------------------------------------------------------------------------- #


def test_item_disimpan_sebagai_salinan(tabel):
    k = kr.Keranjang()
    k.tambah_tabel("Regresi", "Koefisien", tabel)
    tabel.loc[0, "B"] = 999.0
    # Data aktif boleh berubah setelah hasil ditangkap; hasilnya tidak ikut berubah.
    assert k.item[0].tabel.loc[0, "B"] == pytest.approx(0.2423)


def test_isi_yang_sama_persis_ditolak(tabel):
    k = kr.Keranjang()
    assert k.tambah_tabel("Regresi", "Koefisien", tabel) is True
    assert k.tambah_tabel("Regresi", "Koefisien", tabel) is False
    assert len(k.item) == 1


def test_tabel_berbeda_tetap_diterima(tabel):
    k = kr.Keranjang()
    k.tambah_tabel("Regresi", "Koefisien", tabel)
    lain = tabel.copy()
    lain.loc[0, "B"] = 0.5
    assert k.tambah_tabel("Regresi", "Koefisien", lain) is True
    assert len(k.item) == 2


def test_sidik_membedakan_ruas_yang_bersebelahan():
    # Tanpa pemisah, ("ab", "c") dan ("a", "bc") akan bertabrakan.
    a = kr.Item(bagian="ab", judul="c", jenis="tafsiran", teks="t")
    b = kr.Item(bagian="a", judul="bc", jenis="tafsiran", teks="t")
    assert a.sidik != b.sidik


def test_hasil_dari_data_lain_ditandai_basi(isi):
    basi = isi.basi("data#2")
    assert {i.judul for i in basi} == {"Koefisien", "Muatan"}
    # Item tanpa tanda data tidak dianggap basi, karena asalnya tidak diketahui.
    assert isi.basi("data#1") == []


def test_hapus_per_item_dan_per_bagian(isi):
    sidik = isi.item[0].sidik
    assert isi.hapus(sidik) is True
    assert isi.hapus(sidik) is False
    assert isi.hapus_bagian("CFA") == 1
    assert isi.bagian() == ["Regresi linear"]
    isi.kosongkan()
    assert isi.kosong()


def test_ringkas_dan_daftar_isi(isi):
    ringkas = isi.ringkas()
    assert ringkas == {"Bagian": 2, "Tabel": 2, "Tafsiran": 1, "Objek": 3}
    daftar = isi.daftar_isi()
    assert list(daftar["Bagian"]) == ["Regresi linear", "CFA"]
    assert daftar.loc[0, "Jumlah"] == 2


def test_urutan_bagian_mengikuti_penyimpanan_pertama(isi):
    assert isi.bagian() == ["Regresi linear", "CFA"]


@pytest.mark.parametrize(
    "argumen, pesan",
    [
        (dict(bagian="A", judul="B", jenis="grafik"), "tidak dikenal"),
        (dict(bagian="A", judul="B", jenis="tabel"), "memerlukan DataFrame"),
        (dict(bagian="A", judul="B", jenis="tafsiran", teks="   "), "memerlukan teks"),
    ],
)
def test_item_cacat_ditolak(argumen, pesan):
    with pytest.raises(ValueError, match=pesan):
        kr.Item(**argumen)


# --------------------------------------------------------------------------- #
# Ekspor
# --------------------------------------------------------------------------- #


def test_dokumen_dari_keranjang_memuat_seluruh_bagian(isi):
    dokumen = ekspor.dari_keranjang(isi)
    assert dokumen.judul == "Laporan Uji"
    assert "Peneliti Satu" in dokumen.meta
    subjudul = {b.teks for b in dokumen.blok if b.jenis == "subjudul"}
    assert {"Daftar isi", "Regresi linear", "CFA"} <= subjudul
    # Daftar isi ditambah dua tabel hasil.
    assert len(dokumen.tabel()) == 3


def test_keranjang_kosong_ditolak_dengan_pesan_jelas():
    with pytest.raises(ValueError, match="masih kosong"):
        ekspor.dari_keranjang(kr.Keranjang())


@pytest.mark.parametrize("kode", ["docx", "pdf", "xlsx", "pptx", "html", "md", "json", "zip"])
def test_setiap_format_dibangun_dari_keranjang(isi, kode):
    berkas = ekspor.bangun(isi, kode)
    assert isinstance(berkas, bytes) and len(berkas) > 400


def test_docx_dari_keranjang_terbaca_kembali(isi):
    from docx import Document

    dok = Document(io.BytesIO(ekspor.bangun(isi, "docx")))
    teks = "\n".join(p.text for p in dok.paragraphs)
    assert "Laporan Uji" in teks
    assert "Regresi linear" in teks
    assert dok.tables


def test_pdf_dari_keranjang_terbaca_kembali(isi):
    from pypdf import PdfReader

    baca = PdfReader(io.BytesIO(ekspor.bangun(isi, "pdf")))
    assert len(baca.pages) >= 1
    assert "Laporan Uji" in baca.pages[0].extract_text()


def test_xlsx_dari_keranjang_satu_lembar_per_tabel(isi):
    lembar = pd.read_excel(io.BytesIO(ekspor.bangun(isi, "xlsx")), sheet_name=None)
    assert "Keterangan" in lembar
    assert len(lembar) >= 4  # keterangan + daftar isi + dua tabel hasil
    assert all(len(n) <= 31 for n in lembar)


def test_pptx_dari_keranjang_punya_slide_per_bagian(isi):
    from pptx import Presentation

    presentasi = Presentation(io.BytesIO(ekspor.bangun(isi, "pptx")))
    assert len(presentasi.slides) >= 3


def test_json_dari_keranjang_membawa_baris_tabel(isi):
    hasil = json.loads(ekspor.bangun(isi, "json"))
    assert hasil["judul"] == "Laporan Uji"
    tabel = [b for b in hasil["blok"] if b["tabel"]]
    assert tabel and tabel[-1]["tabel"][0]["Butir"] == "x1"


def test_zip_dari_keranjang_tanpa_sintaks(isi):
    """Keranjang tidak membawa konfigurasi analisis, jadi sintaks tidak dibangkitkan."""
    arsip = zipfile.ZipFile(io.BytesIO(ekspor.bangun(isi, "zip")))
    nama = arsip.namelist()
    for akhiran in (".html", ".md", ".docx", ".xlsx", ".pdf", ".json"):
        assert any(n.endswith(akhiran) for n in nama), akhiran
    assert not any(n.startswith("sintaks/") for n in nama)
    assert sum(1 for n in nama if n.startswith("tabel/")) == 3


def test_nama_berkas_keranjang(isi):
    assert ekspor.nama_berkas(isi, "docx").endswith("laporan_hasil.docx")


def test_html_dari_keranjang_memuat_tabel_dan_lolos_escape():
    k = kr.Keranjang()
    k.tambah_tabel(
        "Uji", "Tabel", pd.DataFrame({"Nama": ["<script>alert(1)</script>"], "Nilai": [1]})
    )
    halaman = ekspor.bangun(k, "html").decode("utf-8")
    assert "<table>" in halaman
    # Isi tabel berasal dari data pengguna, jadi harus di-escape.
    assert "<script>alert(1)</script>" not in halaman
    assert "&lt;script&gt;" in halaman
